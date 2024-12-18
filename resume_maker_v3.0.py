import customtkinter as ctk  # For modern GUI design
from tkinter import messagebox  # To show pop-up alerts
import os  # For file and path handling
import pyaudio  # For audio recording/playback
import wave  # To read/write WAV files
import tempfile  # For temporary file creation
import threading  # To run tasks in parallel
import requests  # For API requests
import pyttsx3  # For text-to-speech
import speech_recognition as sr  # For speech-to-text
from PIL import Image  # For image processing
import time  # For time-related functions
import webrtcvad  # For voice activity detection
from fuzzywuzzy import process  # For fuzzy string matching
import json  # For handling JSON data
from dotenv import load_dotenv  # To load environment variables
import docx  # To create/edit Word documents
from docx.shared import Pt, Inches  # To format Word document text and layout
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # To align Word document text
from docx.enum.style import WD_STYLE_TYPE  # To apply Word styles
from docx2pdf import convert  # To convert Word to PDF
from datetime import datetime  # For date/time handling
from tkinter import filedialog  # To open file dialogs
import sys  # For system-specific functions
import traceback  # For error tracebacks
import comtypes.client  # For COM object automation
import winsound  # To play sound or system beeps


# Get absolute path to resource, works for dev and PyInstaller
def resource_path(relative_path):
    try:
        # PyInstaller creates a temporary folder and stores its path in sys._MEIPASS.
        base_path = sys._MEIPASS
    except AttributeError:
        # If not running in a PyInstaller bundle, use the original path.
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# Define the default save directory
DEFAULT_SAVE_DIR = os.path.join(os.path.expanduser("~"), "Downloads", "VoiceResumeMaker")

# Ensure the directory exists
os.makedirs(DEFAULT_SAVE_DIR, exist_ok=True)

# Initialize the save directory
save_directory = DEFAULT_SAVE_DIR

# Define the data subdirectory
data_directory = os.path.join(save_directory, 'data')

# Ensure the data subdirectory exists
os.makedirs(data_directory, exist_ok=True)

# Load environment variables from the .env file using resource_path
dotenv_path = resource_path('.env')
load_dotenv(dotenv_path)

# Set your OpenAI API key from the .env file
openai_api_key = os.getenv('OPENAI_API_KEY')

# Add a function to get the autosave file path
def get_autosave_file_path():
    return os.path.join(data_directory, resource_path('autosave.json'))

# Parses the resume text file and structures the data into a dictionary.
def parse_resume_text(text_file):
    with open(text_file, 'r', encoding='utf-8') as file:
        lines = [line.strip() for line in file.readlines()]

    data = {
        "name": "",
        "phone": "",
        "email": "",
        "address": "",
        "professional_summary": "",
        "skills": [],
        "certifications": [],
        "achievements": [],
        "work_experience": [],
        "education": [],
        "interests": [],
        "extracurricular_activities": [],
        "volunteer_experience": [],
        "professional_associations": [],
        "references": ""
    }

    sections = [
        "Professional Summary", "Skills", "Certifications and Training",
        "Professional Achievements", "Work Experience", "Education",
        "Interests", "Extracurricular Activities", "Volunteer Experience",
        "Professional Associations", "References"
    ]

    current_section = None
    i = 0
    n = len(lines)

    # Parse header (Name, Phone, Email, Address)
    if i < n:
        data["name"] = lines[i]
        i += 1
    if i < n:
        data["phone"] = lines[i]
        i += 1
    if i < n:
        data["email"] = lines[i]
        i += 1
    if i < n:
        data["address"] = lines[i]
        i += 1

    # Parse sections
    while i < n:
        line = lines[i]

        if line in sections:
            current_section = line
            i += 1
            continue

        if not current_section:
            i += 1
            continue

        if current_section == "Professional Summary":
            summary = []
            while i < n and lines[i] not in sections:
                if lines[i]:
                    summary.append(lines[i])
                i += 1
            data["professional_summary"] = ' '.join(summary)
            current_section = None

        elif current_section == "Skills":
            while i < n and lines[i].startswith("- "):
                skill = lines[i][2:].strip()
                data["skills"].append(skill)
                i += 1
            current_section = None

        elif current_section == "Certifications and Training":
            while i < n and lines[i].startswith("- "):
                cert = lines[i][2:].strip()
                data["certifications"].append(cert)
                i += 1
            current_section = None

        elif current_section == "Professional Achievements":
            while i < n and lines[i].startswith("- "):
                achievement = lines[i][2:].strip()
                data["achievements"].append(achievement)
                i += 1
            current_section = None

        elif current_section == "Work Experience":
            while i < n and lines[i] not in sections:
                if not lines[i]:
                    i += 1
                    continue

                # Read company name
                company = lines[i]
                i += 1

                # Initialize position and description
                position = ""
                description = []

                # Check for position
                if i < n and lines[i] and not lines[i] in sections:
                    position = lines[i]
                    i += 1

                # Collect job description
                while i < n and lines[i] not in sections and lines[i]:
                    description.append(lines[i])
                    i += 1

                # Append the job to work_experience
                job = {
                    "company": company,
                    "position": position,
                    "description": ' '.join(description)
                }
                data["work_experience"].append(job)
            current_section = None

        elif current_section == "Education":
            while i < n and lines[i] not in sections:
                if lines[i]:
                    data["education"].append(lines[i])
                i += 1
            current_section = None

        elif current_section == "Interests":
            while i < n and lines[i].startswith("- "):
                interest = lines[i][2:].strip()
                data["interests"].append(interest)
                i += 1
            current_section = None

        elif current_section == "Extracurricular Activities":
            while i < n and lines[i].startswith("- "):
                activity = lines[i][2:].strip()
                data["extracurricular_activities"].append(activity)
                i += 1
            current_section = None

        elif current_section == "Volunteer Experience":
            while i < n and lines[i].startswith("- "):
                volunteer = lines[i][2:].strip()
                data["volunteer_experience"].append(volunteer)
                i += 1
            current_section = None

        elif current_section == "Professional Associations":
            while i < n and lines[i].startswith("- "):
                association = lines[i][2:].strip()
                data["professional_associations"].append(association)
                i += 1
            current_section = None

        elif current_section == "References":
            references = []
            while i < n and lines[i] not in sections:
                if lines[i]:
                    references.append(lines[i])
                i += 1
            data["references"] = ' '.join(references)
            current_section = None

        else:
            i += 1

    return data


def set_styles(doc):
    """
    Defines and sets custom styles for the Word document.
    """
    styles = doc.styles

    # Title Style for Name
    if 'NameStyle' not in styles:
        name_style = styles.add_style('NameStyle', WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.name = 'Arial'
        name_style.font.size = Pt(24)
        name_style.font.bold = True
        name_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name_style.paragraph_format.space_after = Pt(12)

    # Heading Style for Sections
    if 'HeadingStyle' not in styles:
        heading_style = styles.add_style('HeadingStyle', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

    # Subheading Style for Work Experience
    if 'SubheadingStyle' not in styles:
        subheading_style = styles.add_style('SubheadingStyle', WD_STYLE_TYPE.PARAGRAPH)
        subheading_style.font.name = 'Arial'
        subheading_style.font.size = Pt(12)
        subheading_style.font.bold = True
        subheading_style.paragraph_format.space_before = Pt(6)
        subheading_style.paragraph_format.space_after = Pt(3)

    # Normal Text Style
    if 'NormalStyle' not in styles:
        normal_style = styles.add_style('NormalStyle', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.line_spacing = Pt(14)
        normal_style.paragraph_format.space_after = Pt(6)

    # Bullet Point Style
    if 'BulletStyle' not in styles:
        bullet_style = styles.add_style('BulletStyle', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = 'Arial'
        bullet_style.font.size = Pt(11)
        bullet_style.paragraph_format.left_indent = Inches(0.25)
        bullet_style.paragraph_format.space_after = Pt(3)


# Creates a formatted Word document based on the parsed resume data.
def create_resume_document(data, output_docx):
    doc = docx.Document()

    # Set default margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Define styles
    set_styles(doc)

    # Add Name
    name_paragraph = doc.add_paragraph(data["name"], style='NameStyle')

    # Add Contact Information
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_paragraph.style = 'NormalStyle'
    contact_text = f"{data['phone']} | {data['email']} | {data['address']}"
    contact_paragraph.add_run(contact_text).italic = True

    # Add spacing after contact info
    doc.add_paragraph()

    # Add Professional Summary
    if data["professional_summary"]:
        doc.add_paragraph("Professional Summary", style='HeadingStyle')
        summary_paragraph = doc.add_paragraph(data["professional_summary"], style='NormalStyle')

    # Add Skills
    if data["skills"]:
        doc.add_paragraph("Skills", style='HeadingStyle')
        for skill in data["skills"]:
            p = doc.add_paragraph(f"• {skill}", style='BulletStyle')

    # Add Certifications and Training
    if data["certifications"]:
        doc.add_paragraph("Certifications and Training", style='HeadingStyle')
        for cert in data["certifications"]:
            p = doc.add_paragraph(f"• {cert}", style='BulletStyle')

    # Add Professional Achievements
    if data["achievements"]:
        doc.add_paragraph("Professional Achievements", style='HeadingStyle')
        for achievement in data["achievements"]:
            p = doc.add_paragraph(f"• {achievement}", style='BulletStyle')

    # Add Work Experience
    if data["work_experience"]:
        doc.add_paragraph("Work Experience", style='HeadingStyle')
        for job in data["work_experience"]:
            # Company Name
            company_paragraph = doc.add_paragraph(job["company"], style='SubheadingStyle')
            # Position
            if job["position"]:
                position_text = job["position"]
                position_paragraph = doc.add_paragraph(position_text, style='NormalStyle')
                position_paragraph.runs[0].bold = True
            # Description
            if job["description"]:
                description_paragraph = doc.add_paragraph(job["description"], style='NormalStyle')

    # Add Education
    if data["education"]:
        doc.add_paragraph("Education", style='HeadingStyle')
        for edu in data["education"]:
            edu_paragraph = doc.add_paragraph(edu, style='NormalStyle')

    # Add Interests
    if data["interests"]:
        doc.add_paragraph("Interests", style='HeadingStyle')
        for interest in data["interests"]:
            p = doc.add_paragraph(f"• {interest}", style='BulletStyle')

    # Add Extracurricular Activities
    if data["extracurricular_activities"]:
        doc.add_paragraph("Extracurricular Activities", style='HeadingStyle')
        for activity in data["extracurricular_activities"]:
            p = doc.add_paragraph(f"• {activity}", style='BulletStyle')

    # Add Volunteer Experience
    if data["volunteer_experience"]:
        doc.add_paragraph("Volunteer Experience", style='HeadingStyle')
        for volunteer in data["volunteer_experience"]:
            p = doc.add_paragraph(f"• {volunteer}", style='BulletStyle')

    # Add Professional Associations
    if data["professional_associations"]:
        doc.add_paragraph("Professional Associations", style='HeadingStyle')
        for association in data["professional_associations"]:
            p = doc.add_paragraph(f"• {association}", style='BulletStyle')

    # Add References
    if data["references"]:
        doc.add_paragraph("References", style='HeadingStyle')
        references_paragraph = doc.add_paragraph(data["references"], style='NormalStyle')

    # Save the document
    doc.save(output_docx)
      


def convert_to_pdf(word_file, pdf_file):
    try:
        # Create an instance of Word Application
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False

        # Open the Word document
        doc = word.Documents.Open(os.path.abspath(word_file))

        # Save as PDF (FileFormat=17 corresponds to PDF format)
        doc.SaveAs(os.path.abspath(pdf_file), FileFormat=17)

        # Close the document and Word application
        doc.Close()
        word.Quit()
    except Exception as e:
        traceback_str = traceback.format_exc()
        messagebox.showerror("Error", f"Failed to convert to PDF: {e}\n\n{traceback_str}")


def create_finished_resume():
    # Generate unique filenames within data_directory
    filename_formatted_resume = get_unique_filename('Resume', 'txt', directory=data_directory)
    text_file = os.path.join(data_directory, filename_formatted_resume)
    
    # Output documents saved directly into save_directory
    filename_resume_docx = filename_formatted_resume.replace('.txt', '.docx')
    output_docx = os.path.join(save_directory, filename_resume_docx)
    
    filename_resume_pdf = filename_formatted_resume.replace('.txt', '.pdf')
    output_pdf = os.path.join(save_directory, filename_resume_pdf)
    
    # Check if the formatted resume file exists in data_directory
    if not os.path.exists(text_file):
        # Attempt to generate the formatted resume
        formatted_resume = generate_formatted_resume()
        if formatted_resume:
            try:
                with open(text_file, 'w', encoding='utf-8') as f:
                    f.write(formatted_resume)
                # messagebox.showinfo("Info", "Formatted resume has been generated, please wait while the finished resume is generated.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save formatted resume: {e}")
                return
        else:
            messagebox.showerror("Error", "Failed to generate formatted resume via ChatGPT.")
            return

    try:
        # Parse the formatted resume text file from data_directory
        data = parse_resume_text(text_file)

        # Create the Word document in save_directory
        create_resume_document(data, output_docx)

        # Convert to PDF
        if os.path.exists(output_docx):
            convert_to_pdf(output_docx, output_pdf)
            messagebox.showinfo("Success", f"Finished resume has been saved as '{output_docx}' and '{output_pdf}'.")
        else:
            messagebox.showerror("Error", f"Failed to create the Word document '{output_docx}'.")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred while creating the finished resume: {e}")





def clear_data():
    """Clears all transcribed responses and resets the GUI to the initial state."""
    global current_question, transcribed_responses, current_mode

    # Prompt for confirmation
    confirm = messagebox.askyesno("Confirm", "Are you sure you want to clear all data and restart?")
    if not confirm:
        return  # User chose not to clear data

    # Clear the responses and reset the current question
    transcribed_responses.clear()
    current_question = 0
    current_mode = None  # Reset current mode

    # Update the question label to the first question
    if current_mode == 'manual':
        update_question()
    elif current_mode == 'voice':
        update_question_voice_mode(current_question)

    # Reset the transcription label
    transcribed_label.configure(text="Transcription: ")

    # Clear the resume preview textbox
    resume_preview_textbox.configure(state="normal")
    resume_preview_textbox.delete("1.0", "end")
    resume_preview_textbox.configure(state="disabled")

    # Reset the progress bar
    progress_bar.set(0)

    # Optionally, inform the user
    messagebox.showinfo("Data Cleared", "All data has been cleared. You can start fresh.")

    # Update the 'Add More' button state
    if current_mode == 'manual':
        update_add_more_button_state()

    autosave_file = get_autosave_file_path()
    if os.path.exists(autosave_file):
        os.remove(autosave_file)

# Loads questions and their field mappings from a JSON file.
def load_questions(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            questions = [q['question'] for q in data['questions']]
            field_mapping = {q['question']: q['field'] for q in data['questions']}
        return questions, field_mapping
    except FileNotFoundError:
        messagebox.showerror("Error", f"Configuration file '{file_path}' not found.")
        return [], {}
    except json.JSONDecodeError as e:
        messagebox.showerror("Error", f"Error parsing '{file_path}': {e}")
        return [], {}
    except KeyError as e:
        messagebox.showerror("Error", f"Missing key in '{file_path}': {e}")
        return [], {}
    except Exception as e:
        messagebox.showerror("Error", f"Failed to load questions: {e}")
        return [], {}

# Load questions and field mapping
questions, field_mapping = load_questions(resource_path('questions.json'))


# Initialize global variables
current_question = 0
transcribed_responses = {}
add_more_flag = False  # Flag to indicate if we are adding more to the response
voice_thread = None  # Add this at the top with other global variables
current_mode = None  # Track the current mode ('manual' or 'voice')
# Add global variables for the timer
timer_running = False
recording_start_time = None
recording_elapsed_time = 0.0
timer_label = None  # To be initialized in initialize_manual_gui()
start_frame = None  # Add this line
# Global variable to store the selected microphone index
selected_microphone_index = None

# Threading events
stop_recording_event = threading.Event()
pause_recording_event = threading.Event() 

# Function to update the timer label
def update_timer():
    global timer_running, recording_start_time, recording_elapsed_time
    if timer_running:
        elapsed_time = time.time() - recording_start_time + recording_elapsed_time
        formatted_time = time.strftime('%H:%M:%S', time.gmtime(elapsed_time))
        timer_label.configure(text=f"Recording Time: {formatted_time}")
        # Schedule the function to run again after 100 ms
        root.after(100, update_timer)


def load_trinidad_locations(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            # Read all lines and strip any whitespace
            locations = [line.strip() for line in f if line.strip()]
        return locations
    except Exception as e:
        messagebox.showerror("Error", f"Failed to load locations: {e}")

# Load the catalogue of Trinidad locations
places_in_trinidad = load_trinidad_locations(resource_path('trinidad_locations.txt'))


def find_location_matches(transcribed_address, catalogue, threshold=80):
    words = transcribed_address.split()
    matches = []
    max_phrase_length = min(5, len(words))  # Adjust based on the maximum expected length of place names

    # Generate all possible n-grams
    for size in range(1, max_phrase_length + 1):
        for i in range(len(words) - size + 1):
            phrase = ' '.join(words[i:i+size])
            match, score = process.extractOne(phrase, catalogue)
            # Debug statement to see what's being matched
            print(f"Trying phrase: '{phrase}' | Best match: '{match}' with score {score}")
            if score >= threshold:
                matches.append((match, score))

    # Remove duplicates and keep the highest score for each match
    unique_matches = {}
    for match, score in matches:
        if match not in unique_matches or score > unique_matches[match]:
            unique_matches[match] = score

    # Return matches sorted by score
    sorted_matches = sorted(unique_matches.items(), key=lambda x: x[1], reverse=True)
    return [match for match, score in sorted_matches]

def handle_address_transcription(transcribed_text):
    print(f"Transcribed Address: '{transcribed_text}'")  # Debug statement
    matched_locations = find_location_matches(transcribed_text, places_in_trinidad)

    if matched_locations:
        # Remove matched locations from the transcribed text
        other_parts = transcribed_text
        for location in matched_locations:
            other_parts = other_parts.replace(location, '')
        # Construct the final address result
        address_result = f"{other_parts.strip()}, {' '.join(matched_locations)}"
    else:
        address_result = transcribed_text

    print(f"Matched Locations: {matched_locations}")  # Debug statement
    print(f"Final Address Result: '{address_result}'")  # Debug statement

    # Save the address to the responses
    question = "Please provide your address(provide spelling if necessary)."
    transcribed_responses[question] = address_result

    # Update the transcription label
    root.after(0, update_transcription_label, f"Transcription: {address_result}")

# Function to update the status label
def update_status(message):
    if 'status_label' in globals():
        status_label.configure(text=f"Status: {message}")

# Function to handle recording and transcribing for voice mode
def record_and_transcribe_response(record_seconds):
    CHUNK = 1024
    FORMAT = pyaudio.paInt16
    CHANNELS = 1
    RATE = 16000  # Whisper models are trained on 16kHz audio

    p = pyaudio.PyAudio()
    frames = []

    try:
        stream = p.open(format=FORMAT,
                        channels=CHANNELS,
                        rate=RATE,
                        input=True,
                        frames_per_buffer=CHUNK)

        # Inform the user to start speaking
        root.after(0, update_status, "Recording... Please speak now.")

        for _ in range(0, int(RATE / CHUNK * record_seconds)):
            data = stream.read(CHUNK)
            frames.append(data)

        # Finish recording
        stream.stop_stream()
        stream.close()
        p.terminate()

        root.after(0, update_status, "Transcribing response...")

        # Save the recorded data as a WAV file
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp_file:
            wf = wave.open(tmp_file.name, 'wb')
            wf.setnchannels(CHANNELS)
            wf.setsampwidth(p.get_sample_size(FORMAT))
            wf.setframerate(RATE)
            wf.writeframes(b''.join(frames))
            wf.close()
            temp_filename = tmp_file.name

        # Retry mechanism for the API call
        max_retries = 5
        retry_count = 0
        backoff_time = 2  # Start with a 2-second backoff

        while retry_count < max_retries:
            try:
                api_url = 'https://api.openai.com/v1/audio/transcriptions'
                headers = {
                    'Authorization': f'Bearer {openai_api_key}',
                }

                with open(temp_filename, 'rb') as audio_file:
                    files = {
                        'file': (temp_filename, audio_file, 'audio/wav'),
                        'model': (None, 'whisper-1'),
                    }
                    data = {
                        'language': selected_language_code  # Add language parameter
                    }

                    response = requests.post(api_url, headers=headers, files=files, data=data)
                    response.raise_for_status()  # Raise exception for bad responses

                    result = response.json()
                    transcription_text = result['text'].strip()
                    break  # Break out of the retry loop if the request is successful

            except (requests.ConnectionError, requests.Timeout) as e:
                retry_count += 1
                if retry_count == max_retries:
                    messagebox.showerror("Error", f"Failed to connect after {max_retries} attempts. Please check your internet connection.")
                    return ""
                else:
                    root.after(0, update_status, f"Connection error, retrying in {backoff_time} seconds...")
                    time.sleep(backoff_time)  # Wait before retrying
                    backoff_time *= 2  # Exponential backoff
            except requests.HTTPError as e:
                messagebox.showerror("Error", f"HTTP Error: {e}")
                return ""
            except Exception as e:
                messagebox.showerror("Error", f"Unexpected error: {e}")
                return ""
            finally:
                os.unlink(temp_filename)  # Remove the temp file

        return transcription_text

    except Exception as e:
        messagebox.showerror("Error", f"Error during recording: {e}")
        return ""


def listen_and_transcribe(recognizer, microphone):
    try:
        stop_phrase = "stop recording"
        transcription = []

        with microphone as source:
            recognizer.adjust_for_ambient_noise(source)
            root.after(0, update_status, "Listening...")

            while True:
                audio = recognizer.listen(source, phrase_time_limit=5)

                try:
                    # Recognize speech using Google's speech recognition
                    text = recognizer.recognize_google(audio, language=selected_language_code)
                    print(f"You said: {text}")

                    if stop_phrase.lower() in text.lower():
                        print("Stop phrase detected. Ending recording.")
                        break
                    else:
                        transcription.append(text)

                except sr.UnknownValueError:
                    print("Could not understand audio. Please try again.")
                except sr.RequestError as e:
                    messagebox.showerror("Error", f"Could not request results from Google Speech Recognition service; {e}")
                    return None

        full_transcription = ' '.join(transcription)
        return full_transcription.strip()

    except Exception as e:
        messagebox.showerror("Error", f"Error during listening: {e}")
        return None

# Function to update the question label in voice mode
def update_question_voice_mode(idx):
    question = questions[idx]
    question_label.configure(text=f"Question {idx + 1}/{len(questions)}:\n{question}")
    # Update progress bar
    progress = (idx + 1) / len(questions)
    progress_bar.set(progress)
    # Update transcription label if responses exist
    if question in transcribed_responses and transcribed_responses[question]:
        transcribed_label.configure(text=f"Transcription: {transcribed_responses[question]}")
    else:
        transcribed_label.configure(text="Transcription: ")
    # Update resume preview
    update_resume_preview()

# Function to update the transcription label
def update_transcription_label(transcription_text):
    transcribed_label.configure(text=f"Transcription: {transcription_text}")
    update_resume_preview()

def listen_for_command(recognizer, microphone):
    try:
        with microphone as source:
            recognizer.adjust_for_ambient_noise(source)
            print("Listening for command...")
            audio = recognizer.listen(source, timeout=5, phrase_time_limit=5)
        # Recognize speech using Google's speech recognition
        command = recognizer.recognize_google(audio, language=selected_language_code)
        print(f"Command recognized: {command}")
        return command.lower()
    except sr.UnknownValueError:
        engine.say("Sorry, I didn't catch that.")
        engine.runAndWait()
        return None
    except sr.WaitTimeoutError:
        engine.say("No speech detected.")
        engine.runAndWait()
        return None
    except sr.RequestError as e:
        messagebox.showerror("Error", f"Could not request results from Google Speech Recognition service; {e}")
        return None



def listen_and_transcribe_with_stop_phrase(recognizer, microphone):
    CHUNK_DURATION_MS = 30  # Duration of a chunk in milliseconds
    SAMPLE_RATE = 16000  # Sample rate in Hz
    CHUNK_SIZE = int(SAMPLE_RATE * CHUNK_DURATION_MS / 1000)  # Chunk size in samples
    FORMAT = pyaudio.paInt16  # 16-bit int sampling
    CHANNELS = 1  # Mono audio

    # Initialize VAD
    vad = webrtcvad.Vad(1)  # Aggressiveness from 0 to 3
    frames = []
    silence_threshold = 4  # Seconds of silence before stopping
    silence_counter = 0  # Counter for silence duration

    p = pyaudio.PyAudio()
    try:
        if selected_microphone_index is not None:
            stream = p.open(format=FORMAT,
                            channels=CHANNELS,
                            rate=SAMPLE_RATE,
                            input=True,
                            input_device_index=selected_microphone_index,
                            frames_per_buffer=CHUNK_SIZE)
        else:
            stream = p.open(format=FORMAT,
                            channels=CHANNELS,
                            rate=SAMPLE_RATE,
                            input=True,
                            frames_per_buffer=CHUNK_SIZE)
    except Exception as e:
        messagebox.showerror("Error", f"Could not open microphone: {e}")
        return None

    root.after(0, update_status, "Listening... Please speak now.")
    print("Listening... Please speak now.")

    try:
        while True:
            frame = stream.read(CHUNK_SIZE, exception_on_overflow=False)
            is_speech = vad.is_speech(frame, SAMPLE_RATE)

            if is_speech:
                frames.append(frame)
                silence_counter = 0  # Reset silence counter when speech is detected
            else:
                silence_counter += CHUNK_DURATION_MS / 1000.0  # Increment silence counter
                if silence_counter >= silence_threshold:
                    print("Silence detected. Stopping recording.")
                    break
                frames.append(frame)  # Optionally include silence frames
    except Exception as e:
        messagebox.showerror("Error", f"Error during recording: {e}")
    finally:
        stream.stop_stream()
        stream.close()
        p.terminate()

    root.after(0, update_status, "Transcribing response...")
    # Save the recorded data as a WAV file
    temp_filename = None
    try:
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp_file:
            wf = wave.open(tmp_file.name, 'wb')
            wf.setnchannels(CHANNELS)
            wf.setsampwidth(p.get_sample_size(FORMAT))
            wf.setframerate(SAMPLE_RATE)
            wf.writeframes(b''.join(frames))
            wf.close()
            temp_filename = tmp_file.name

        # Use requests to call the OpenAI API for transcription
        try:
            api_url = 'https://api.openai.com/v1/audio/transcriptions'

            headers = {
                'Authorization': f'Bearer {openai_api_key}',
            }

            with open(temp_filename, 'rb') as audio_file:
                files = {
                    'file': (temp_filename, audio_file, 'audio/wav'),
                }
                data = {
                    'model': 'whisper-1',
                    'language': selected_language_code  # Add language parameter
                }

                response = requests.post(api_url, headers=headers, files=files, data=data)
                response.raise_for_status()

                result = response.json()
                transcription_text = result['text'].strip()
        except Exception as e:
            error_message = ""
            try:
                error_message = response.json().get('error', {}).get('message', str(e))
            except:
                error_message = str(e)
            messagebox.showerror("Error", f"Transcription failed: {error_message}")
            transcription_text = ""
    finally:
        if temp_filename:
            os.unlink(temp_filename)  # Remove the temp file

    return transcription_text

# Function to reset the voice walkthrough
def reset_voice_walkthrough():
    """Resets the voice walkthrough by clearing responses and resetting the question index."""
    global transcribed_responses, current_question
    transcribed_responses.clear()
    current_question = 0
    update_question_voice_mode(current_question)
    update_transcription_label("Transcription: ")
    update_resume_preview()
    print("Voice walkthrough has been reset.")  # Debug statement

# Function to play the beep sound
def play_beep_sound():
    def play_sound():
        try:
            winsound.PlaySound(resource_path('beep.wav'), winsound.SND_FILENAME | winsound.SND_ASYNC)
        except Exception as e:
            print(f"Error playing beep sound: {e}")
    threading.Thread(target=play_sound).start()

# Helper function for assistant speech and status updates
def assistant_say(message, update_status_label=True):
    if update_status_label:
        root.after(0, update_status, message)
    engine.say(message)
    engine.runAndWait()

# Function to perform the voice-assisted walkthrough
def voice_walkthrough():
    global transcribed_responses, engine, current_question
    engine = pyttsx3.init()
    engine.setProperty('rate', tts_settings['rate'])  # Use rate from settings
    engine.setProperty('voice', tts_settings['voice'].id)  # Use selected voice from settings

    while True:  # Outer loop to allow restarting the walkthrough
        # Provide initial instructions once
        initial_instructions = (
            "Welcome to the voice-assisted resume generator. "
            "You will be asked a series of questions to build your resume. "
            "After each question, please provide your response. "
            "When you are finished speaking, simply stop talking, and I will process your response. "
            "After your response, I will read it back to you. "
            "If it is correct, please say 'yes'. "
            "If it is incorrect, please say 'no' to redo your response. "
            "If you wish to restart the voice walkthrough at any time, say 'restart voice walkthrough'. "
            "Let's begin."
        )
        assistant_say(initial_instructions)  # Use assistant_say instead of engine.say()

        recognizer = sr.Recognizer()
        microphone = sr.Microphone()

        for idx in range(current_question, len(questions)):
            if stop_recording_event.is_set():
                return  # Exit if a stop is requested

            response_recorded = False
            while not response_recorded:
                current_question = idx  # Update current_question
                # Update GUI with current question
                root.after(0, update_question_voice_mode, idx)
                # Update status
                root.after(0, update_status, f"Asking Question {idx+1}/{len(questions)}")

                # Speak the question
                assistant_say(questions[idx])  # Use assistant_say

                # Indicate that recording is about to start
                assistant_say("Now recording your response, please speak after the beep.")

                # Play the beep sound
                play_beep_sound()

                # Record and transcribe the user's response
                transcription_text = listen_and_transcribe_with_stop_phrase(recognizer, microphone)

                if transcription_text is None or transcription_text.strip() == "":
                    assistant_say("I didn't catch that. Let's try again.")
                    continue  # Restart the question

                # Check for restart phrase
                if "restart voice walkthrough" in transcription_text.lower():
                    assistant_say("Restarting the voice walkthrough. All previous responses have been cleared.")
                    # Halt current operations by setting the stop event
                    stop_recording_event.set()
                    # Reset the walkthrough
                    reset_voice_walkthrough()
                    # Clear the stop event for future recordings
                    stop_recording_event.clear()
                    # Break out of the current for-loop to restart
                    break

                # Inform the user that processing is underway
                assistant_say("I heard you, processing now.")

                # Inform that transcription is in progress
                assistant_say("Transcribing your response.")

                # Update the status label
                root.after(0, update_status, "Transcribing your response...")

                # Special handling for the address question
                if questions[idx] == "Please provide your address(provide spelling if necessary).":
                    handle_address_transcription(transcription_text)
                    address_result = transcribed_responses[questions[idx]]
                    # Read back the address result
                    assistant_say("You said:")
                    assistant_say(address_result)
                else:
                    # Standard processing for other questions
                    transcribed_responses[questions[idx]] = transcription_text
                    # Update the transcription label
                    root.after(0, update_transcription_label, transcription_text)
                    # Read back the response
                    assistant_say("You said:")
                    assistant_say(transcription_text)

                # Confirmation loop
                confirmation_received = False
                while not confirmation_received:
                    # Ask if the response is correct
                    assistant_say("Is this correct? Please say 'yes' to confirm or 'no' to try again.")

                    # Listen for confirmation
                    confirmation = listen_for_command(recognizer, microphone)
                    if confirmation is not None and 'yes' in confirmation.lower():
                        # Save the response (already saved for the address question)
                        if questions[idx] != "Please provide your address(provide spelling if necessary).":
                            transcribed_responses[questions[idx]] = transcription_text
                        response_recorded = True  # Move to the next question
                        confirmation_received = True  # Exit confirmation loop
                        # Update the resume preview
                        root.after(0, update_resume_preview)
                    elif confirmation is not None and 'no' in confirmation.lower():
                        assistant_say("Let's try again.")
                        confirmation_received = True  # Exit confirmation loop to redo the question
                        # Do not set response_recorded to True, so the question will be re-asked
                    else:
                        assistant_say("I didn't understand your response. Please say 'yes' or 'no'.")
                        # Continue the confirmation loop without changing confirmation_received

            else:
                continue  # Continue to the next question

            # If 'restart voice walkthrough' was detected, restart the outer loop
            if "restart voice walkthrough" in transcription_text.lower():
                break  # Breaks the for-loop and restarts the while-loop

        else:
            # Completed all questions without a restart
            break  # Exit the while-loop

    # After all questions are done
    root.after(0, update_status, "Voice Walkthrough Completed")
    # Generate and save the formatted resume
    create_finished_resume()
    # Inform the user
    assistant_say("Voice walkthrough completed. Your formatted resume has been generated and saved to your selected file location, or if none has been selected it is saved in your downloads folder.")



# Function to record and transcribe using OpenAI Whisper API
def listen_and_transcribe_with_whisper():
    CHUNK_DURATION_MS = 30  # Duration of a chunk in milliseconds
    SAMPLE_RATE = 16000  # Sample rate in Hz
    CHUNK_SIZE = int(SAMPLE_RATE * CHUNK_DURATION_MS / 1000)  # Chunk size in samples
    FORMAT = pyaudio.paInt16  # 16-bit int sampling
    CHANNELS = 1  # Mono audio

    # Initialize VAD
    vad = webrtcvad.Vad(1)  # Aggressiveness from 0 to 3
    frames = []
    silence_threshold = 4  # Seconds of silence before stopping
    silence_counter = 0  # Counter for silence duration

    p = pyaudio.PyAudio()
    stream = p.open(format=FORMAT,
                    channels=CHANNELS,
                    rate=SAMPLE_RATE,
                    input=True,
                    frames_per_buffer=CHUNK_SIZE)

    root.after(0, update_status, "Listening... Please speak now.")
    print("Listening... Please speak now.")

    try:
        while True:
            frame = stream.read(CHUNK_SIZE, exception_on_overflow=False)
            is_speech = vad.is_speech(frame, SAMPLE_RATE)

            if is_speech:
                frames.append(frame)
                silence_counter = 0  # Reset silence counter when speech is detected
            else:
                silence_counter += CHUNK_DURATION_MS / 1000.0  # Increment silence counter
                if silence_counter >= silence_threshold:
                    print("Silence detected. Stopping recording.")
                    break
                frames.append(frame)  # Optionally include silence frames
    except Exception as e:
        messagebox.showerror("Error", f"Error during recording: {e}")
    finally:
        stream.stop_stream()
        stream.close()
        p.terminate()

    root.after(0, update_status, "Transcribing response...")
    # Save the recorded data as a WAV file
    temp_filename = None
    try:
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp_file:
            wf = wave.open(tmp_file.name, 'wb')
            wf.setnchannels(CHANNELS)
            wf.setsampwidth(p.get_sample_size(FORMAT))
            wf.setframerate(SAMPLE_RATE)
            wf.writeframes(b''.join(frames))
            wf.close()
            temp_filename = tmp_file.name

        # Use requests to call the OpenAI API for transcription
        try:
            api_url = 'https://api.openai.com/v1/audio/transcriptions'

            headers = {
                'Authorization': f'Bearer {openai_api_key}',
            }

            with open(temp_filename, 'rb') as audio_file:
                files = {
                    'file': (temp_filename, audio_file, 'audio/wav'),
                    'model': (None, 'whisper-1'),
                }
                data = {
                    'language': selected_language_code  # Add language parameter
                }

                response = requests.post(api_url, headers=headers, files=files, data=data)
                response.raise_for_status()

                result = response.json()
                transcription_text = result['text'].strip()
        except Exception as e:
            error_message = response.json().get('error', {}).get('message', str(e))
            messagebox.showerror("Error", f"Transcription failed: {error_message}")
            transcription_text = ""
    finally:
        if temp_filename:
            os.unlink(temp_filename)  # Remove the temp file

    return transcription_text

# Function to handle recording and transcribing
def recording_thread_function():
    global transcribed_responses, add_more_flag

    CHUNK = 1024
    FORMAT = pyaudio.paInt16
    CHANNELS = 1
    RATE = 16000  # Whisper models are trained on 16kHz audio

    p = pyaudio.PyAudio()

    try:
        if selected_microphone_index is not None:
            stream = p.open(format=FORMAT,
                            channels=CHANNELS,
                            rate=RATE,
                            input=True,
                            input_device_index=selected_microphone_index,
                            frames_per_buffer=CHUNK)
        else:
            stream = p.open(format=FORMAT,
                            channels=CHANNELS,
                            rate=RATE,
                            input=True,
                            frames_per_buffer=CHUNK)
    except Exception as e:
        messagebox.showerror("Error", f"Could not open microphone: {e}")
        root.after(0, update_status, "Ready")
        return False

    frames = []

    print("Recording...")
    root.after(0, update_status, "Recording...")

    while not stop_recording_event.is_set():
        if pause_recording_event.is_set():
            # If paused, wait until the pause is lifted
            time.sleep(0.1)
            continue
        try:
            data = stream.read(CHUNK)
            frames.append(data)
        except Exception as e:
            messagebox.showerror("Error", f"Error while recording: {e}")
            break

    print("Finished recording.")
    root.after(0, update_status, "Processing transcription...")

    stream.stop_stream()
    stream.close()
    p.terminate()

    # Save the recorded data as a WAV file
    with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp_file:
        wf = wave.open(tmp_file.name, 'wb')
        wf.setnchannels(CHANNELS)
        wf.setsampwidth(p.get_sample_size(FORMAT))
        wf.setframerate(RATE)
        wf.writeframes(b''.join(frames))
        wf.close()
        temp_filename = tmp_file.name

    # Use requests to call the OpenAI API for transcription
    try:
        api_url = 'https://api.openai.com/v1/audio/transcriptions'

        headers = {
            'Authorization': f'Bearer {openai_api_key}',
        }

        with open(temp_filename, 'rb') as audio_file:
            files = {
                'file': (temp_filename, audio_file, 'audio/wav'),
            }
            data = {
                'model': 'whisper-1',
                'language': selected_language_code  # Add language parameter
            }

            response = requests.post(api_url, headers=headers, files=files, data=data)
            response.raise_for_status()

            result = response.json()
            transcription_text = result['text'].strip()
    except Exception as e:
        error_message = ""
        try:
            error_message = response.json().get('error', {}).get('message', str(e))
        except:
            error_message = str(e)
        messagebox.showerror("Error", f"Transcription failed: {error_message}")
        transcription_text = ""
    finally:
        os.unlink(temp_filename)  # Remove the temp file

    # Update the transcription label
    root.after(0, update_transcription_label, transcription_text)
    # Update 'Add More' button state
    root.after(0, update_add_more_button_state)

    # Save responses
    current_q = questions[current_question]
    if add_more_flag:
        # Append to existing responses
        transcribed_responses[current_q] += " " + transcription_text if current_q in transcribed_responses else transcription_text
    else:
        # Replace existing responses
        transcribed_responses[current_q] = transcription_text

    # Reset add_more_flag
    add_more_flag = False

    # Schedule GUI updates
    root.after(0, update_resume_preview)
    root.after(0, enable_buttons_after_recording)
    root.after(0, update_status, "Ready")
    root.after(0, show_decision_frame)




# Function to start recording in a new thread
def start_recording_thread(add_more=False):
    global add_more_flag, timer_running, recording_start_time, recording_elapsed_time
    stop_recording_event.clear()
    pause_recording_event.clear()  # Clear pause at start
    add_more_flag = add_more
    # Hide decision frame first
    hide_decision_frame()
    # Disable buttons during recording
    disable_buttons_for_recording()
    # Enable the Pause button with initial text and color
    pause_button.configure(state="normal", text="Pause Recording", fg_color="green")  # Set to green indicating active recording
    # Update labels to indicate recording
    transcribed_label.configure(text="Transcription: Recording...")
    update_status("Recording...")
    # Reset timer variables
    recording_start_time = time.time()
    recording_elapsed_time = 0.0
    timer_running = True
    update_timer()
    # Start recording thread
    recording_thread = threading.Thread(target=recording_thread_function)
    recording_thread.start()


# Function to stop recording
def stop_recording():
    global timer_running, recording_start_time, recording_elapsed_time
    stop_recording_event.set()
    # Reset the pause event in case recording was paused
    pause_recording_event.clear()
    # Update label to indicate processing
    transcribed_label.configure(text="Transcription: Processing...")
    update_status("Processing transcription...")
    # Disable the stop and pause buttons to prevent multiple presses
    stop_button.configure(state="disabled")
    pause_button.configure(state="disabled", text="Pause Recording", fg_color="green")  # Reset to initial state
    # Stop the timer
    timer_running = False
    # Reset timer variables
    recording_start_time = None
    recording_elapsed_time = 0.0
    timer_label.configure(text="Recording Time: 00:00:00")  # Reset the timer label

# Function to toggle pause/resume recording
def toggle_pause_recording():
    global timer_running, recording_start_time, recording_elapsed_time
    if not pause_recording_event.is_set():
        # Pause the recording
        pause_recording_event.set()
        pause_button.configure(text="Resume Recording", fg_color="red")  # Change text and color to red
        update_status("Recording Paused.")
        # Pause the timer
        timer_running = False
        # Calculate elapsed time so far
        recording_elapsed_time += time.time() - recording_start_time
    else:
        # Resume the recording
        pause_recording_event.clear()
        pause_button.configure(text="Pause Recording", fg_color="green")  # Revert text and color to green
        update_status("Recording Resumed.")
        # Resume the timer
        recording_start_time = time.time()
        timer_running = True
        update_timer()



# Function to disable buttons during recording
def disable_buttons_for_recording():
    start_button.configure(state="disabled")
    stop_button.configure(state="normal")
    pause_button.configure(state="normal")  # Enable Pause Button
    next_button.configure(state="disabled")
    prev_button.configure(state="disabled")
    save_button.configure(state="disabled")
    skip_button.configure(state="disabled")
    add_more_main_button.configure(state="disabled")  # Disable 'Add More' button

# Function to re-enable buttons after recording
def enable_buttons_after_recording():
    start_button.configure(state="normal")
    stop_button.configure(state="disabled")
    pause_button.configure(state="disabled", text="Pause Recording", fg_color="green")  # Reset to initial state
    next_button.configure(state="normal")
    prev_button.configure(state="normal")
    save_button.configure(state="normal")
    skip_button.configure(state="normal")
    update_add_more_button_state()  # Update 'Add More' button state


# Function to navigate to the next question
def next_question_func():
    global current_question
    if current_question < len(questions) - 1:
        current_question += 1
        update_question()
    else:
        messagebox.showinfo("End", "You have reached the last question.")

# Function to navigate to the previous question
def prev_question_func():
    global current_question
    if current_question > 0:
        current_question -= 1
        update_question()
    else:
        messagebox.showinfo("Start", "You are at the first question.")

# Function to update the question label and display existing responses if any
def update_question():
    question = questions[current_question]
    question_label.configure(text=f"Question {current_question + 1}/{len(questions)}:\n{question}")

    # Update transcription label if responses exist
    if question in transcribed_responses and transcribed_responses[question]:
        transcribed_label.configure(text=f"Transcription: {transcribed_responses[question]}")
    else:
        transcribed_label.configure(text="Transcription: ")

    # Update 'Add More' button state
    update_add_more_button_state()

    # Update progress bar
    progress = (current_question + 1) / len(questions)
    progress_bar.set(progress)


# Function to generate resume text from responses
def generate_resume_text():
    # If there are no responses, return an empty string
    if not any(transcribed_responses.values()):
        return ""

    # Define the resume sections based on the provided prompt
    resume_sections = {
        "Name": transcribed_responses.get(questions[0], "").strip(),
        "Contact Number": transcribed_responses.get(questions[1], "").strip(),
        "Email Address": transcribed_responses.get(questions[2], "").strip(),
        "Address": transcribed_responses.get(questions[3], "").strip(),
        "Professional Summary": transcribed_responses.get(questions[4], "").strip(),
        "Skills": transcribed_responses.get(questions[5], "").strip(),
        "Certifications and Training": transcribed_responses.get(questions[6], "").strip(),
        "Professional Achievements": transcribed_responses.get(questions[7], "").strip(),
        "Interests": transcribed_responses.get(questions[13], "").strip(),
        "Extracurricular Activities": transcribed_responses.get(questions[14], "").strip(),
        "Volunteer Experience": transcribed_responses.get(questions[15], "").strip(),
        "Professional Associations": transcribed_responses.get(questions[16], "").strip(),
        "References": "Available upon request."
    }

    # Start building the resume text
    resume_text = ""
    significant_sections_added = 0  # Counter for significant sections added

    # Function to add section if response exists
    def add_section(title, content, is_significant=False):
        nonlocal resume_text, significant_sections_added
        if content.strip():
            resume_text += f"{title}:\n{content}\n\n"
            if is_significant:
                significant_sections_added += 1

    # Add Name section separately
    if resume_sections["Name"]:
        resume_text += f"Name:\n{resume_sections['Name']}\n\n"

    # Add Contact Information
    if resume_sections["Contact Number"]:
        resume_text += f"Contact Number:\n{resume_sections['Contact Number']}\n\n"
    if resume_sections["Email Address"]:
        resume_text += f"Email Address:\n{resume_sections['Email Address']}\n\n"
    if resume_sections["Address"]:
        resume_text += f"Address:\n{resume_sections['Address']}\n\n"

    # Add significant sections
    add_section("Professional Summary", resume_sections["Professional Summary"], is_significant=True)

    # Handle Skills
    if resume_sections["Skills"]:
        skills_list = [skill.strip() for skill in resume_sections["Skills"].split(',') if skill.strip()]
        if skills_list:
            resume_text += "Skills:\n"
            for skill in skills_list:
                resume_text += f"- {skill}\n"
            resume_text += "\n"
            significant_sections_added += 1

    # Handle Certifications and Training
    if resume_sections["Certifications and Training"]:
        certs_list = [cert.strip() for cert in resume_sections["Certifications and Training"].split(',') if cert.strip()]
        if certs_list:
            resume_text += "Certifications and Training:\n"
            for cert in certs_list:
                resume_text += f"- {cert}\n"
            resume_text += "\n"
            significant_sections_added += 1

    # Handle Professional Achievements
    if resume_sections["Professional Achievements"]:
        achievements_list = [ach.strip() for ach in resume_sections["Professional Achievements"].split(',') if ach.strip()]
        if achievements_list:
            resume_text += "Professional Achievements:\n"
            for ach in achievements_list:
                resume_text += f"- {ach}\n"
            resume_text += "\n"
            significant_sections_added += 1

    # Handle Work Experience
    work_place = transcribed_responses.get(questions[8], '').strip()
    position_description = transcribed_responses.get(questions[9], '').strip()
    key_projects = transcribed_responses.get(questions[10], '').strip()

    work_experience_content = ""
    if work_place:
        work_experience_content += f"{work_place}\n"
    if position_description:
        work_experience_content += f"{position_description}\n"
    if key_projects:
        work_experience_content += f"Key Projects:\n{key_projects}\n"
    if work_experience_content.strip():
        resume_text += f"Work Experience:\n{work_experience_content.strip()}\n\n"
        significant_sections_added += 1

    # Handle Education
    prior_education = transcribed_responses.get(questions[11], '').strip()
    current_education = transcribed_responses.get(questions[12], '').strip()
    education_content = ""
    if prior_education:
        education_content += f"{prior_education}\n"
    if current_education:
        education_content += f"{current_education}\n"
    if education_content.strip():
        resume_text += f"Education:\n{education_content.strip()}\n\n"
        significant_sections_added += 1

    # Handle Interests
    if resume_sections["Interests"]:
        interests_list = [interest.strip() for interest in resume_sections["Interests"].split(',') if interest.strip()]
        if interests_list:
            resume_text += "Interests:\n"
            for interest in interests_list:
                resume_text += f"- {interest}\n"
            resume_text += "\n"
            significant_sections_added += 1

    # Handle Extracurricular Activities
    add_section("Extracurricular Activities", resume_sections["Extracurricular Activities"], is_significant=True)

    # Handle Volunteer Experience
    add_section("Volunteer Experience", resume_sections["Volunteer Experience"], is_significant=True)

    # Handle Professional Associations
    if resume_sections["Professional Associations"]:
        associations_list = [assoc.strip() for assoc in resume_sections["Professional Associations"].split(',') if assoc.strip()]
        if associations_list:
            resume_text += "Professional Associations:\n"
            for assoc in associations_list:
                resume_text += f"- {assoc}\n"
            resume_text += "\n"
            significant_sections_added += 1

    # Only add References if significant sections were added
    if significant_sections_added > 0:
        resume_text += f"References:\n{resume_sections['References']}\n"

    return resume_text.strip()

# Function to update the resume preview textbox
def update_resume_preview():
    if 'resume_preview_textbox' in globals():
        resume_text = generate_resume_text()
        resume_preview_textbox.configure(state="normal")
        resume_preview_textbox.delete("1.0", "end")
        resume_preview_textbox.insert("end", resume_text)
        resume_preview_textbox.configure(state="disabled")  # Make it read-only
        # Save responses to file
        save_responses_to_file()
    else:
        print("Resume preview textbox is not initialized yet.")


# Function to save the responses to a text file without replacing field names
def save_responses_to_file():
    try:
        filename = get_unique_filename('Responses', 'txt', directory=data_directory)
        filepath = os.path.join(data_directory, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            for question in questions:
                field_name = field_mapping.get(question, question)
                response = transcribed_responses.get(question, "")
                f.write(f"{field_name} - {response}\n")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save responses: {e}")



# Function to save the resume to a text file and generate formatted resume via ChatGPT
def save_resume():
    resume_text = generate_resume_text()
    try:
        filename = get_unique_filename('Resume', 'txt', directory=data_directory)
        filepath = os.path.join(data_directory, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(resume_text)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save basic resume: {e}")
        return

    # Generate formatted resume via ChatGPT
    formatted_resume = generate_formatted_resume()
    if formatted_resume:
        try:
            filename_formatted = get_unique_filename('Formatted_Resume', 'txt', directory=data_directory)
            filepath_formatted = os.path.join(data_directory, filename_formatted)
            with open(filepath_formatted, 'w', encoding='utf-8') as f:
                f.write(formatted_resume)
            messagebox.showinfo("Success", f"Your formatted resume has been saved as '{filename_formatted}'.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save formatted resume: {e}")
    else:
        messagebox.showerror("Error", "Failed to generate formatted resume via ChatGPT.")



# Generates a formatted resume by loading the resume prompt template, inserting the user responses into the Input section, and sending it to ChatGPT.
def generate_formatted_resume():
    try:
        # Load the resume prompt template
        with open(resource_path('resume_prompt.txt'), 'r', encoding='utf-8') as f:
            resume_template = f.read()

        # Construct the input_section with [FIELD] - value pairs
        input_section = ""
        for question in questions:
            field = field_mapping.get(question, question)  # Get the field name
            response = transcribed_responses.get(question, "").strip()
            if response:  # Only include fields that have responses
                input_section += f"{field} - {response}\n"

        # Replace the {input_section} placeholder in the template
        filled_resume_template = resume_template.replace("{input_section}", input_section.strip())

        # Prepare the full prompt for OpenAI
        full_prompt = filled_resume_template

        # Send the prompt to the OpenAI API
        api_url = 'https://api.openai.com/v1/chat/completions'

        headers = {
            'Authorization': f'Bearer {openai_api_key}',
            'Content-Type': 'application/json'
        }

        data = {
            "model": "gpt-3.5-turbo",  # You can use a different model if preferred
            "messages": [
                {"role": "user", "content": full_prompt}
            ],
            "max_tokens": 2000,  # Adjust as needed
            "n": 1,
            "temperature": 0.7
        }

        response = requests.post(api_url, headers=headers, json=data)
        response.raise_for_status()

        result = response.json()
        formatted_resume = result['choices'][0]['message']['content'].strip()
        
        return formatted_resume

    except FileNotFoundError:
        messagebox.showerror("Error", "The resume prompt file 'resume_prompt.txt' was not found.")
        return None
    except requests.exceptions.HTTPError as e:
        error_message = response.json().get('error', {}).get('message', str(e))
        messagebox.showerror("Error", f"OpenAI API request failed: {error_message}")
        return None
    except Exception as e:
        messagebox.showerror("Error", f"Failed to generate formatted resume: {e}")
        return None



# Function to generate and save the formatted resume via ChatGPT
def generate_and_save_formatted_resume():
    update_status("Generating formatted resume...")
    try:
        formatted_resume = generate_formatted_resume()
        if formatted_resume:
            try:
                # Save the formatted resume .txt file in data_directory
                filename_formatted = get_unique_filename('Formatted_Resume', 'txt', directory=data_directory)
                filepath_formatted = os.path.join(data_directory, filename_formatted)
                with open(filepath_formatted, 'w', encoding='utf-8') as f:
                    f.write(formatted_resume)
                messagebox.showinfo("Success", f"Your formatted resume has been saved as '{filename_formatted}'.")
                
                # Proceed to create the final .docx and .pdf in save_directory
                create_finished_resume()    
                
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save formatted resume: {e}")
        else:
            messagebox.showerror("Error", "Failed to generate formatted resume via ChatGPT.")
    finally:
        update_status("Ready")



# Function to skip the current question
def skip_question_func():
    current_q = questions[current_question]
    transcribed_responses[current_q] = ""
    update_resume_preview()
    update_add_more_button_state()
    next_question_func()


# Function to show decision frame
def show_decision_frame():
    decision_frame.pack(pady=10, padx=20, fill="x")
    # Disable main buttons
    start_button.configure(state="disabled")
    stop_button.configure(state="disabled")
    pause_button.configure(state="disabled")  # --- New Feature: Disable Pause Button ---
    next_button.configure(state="disabled")
    prev_button.configure(state="disabled")
    save_button.configure(state="disabled")
    skip_button.configure(state="disabled")
    add_more_main_button.configure(state="disabled")  # Disable 'Add More' button

# Function to hide decision frame
def hide_decision_frame():
    decision_frame.pack_forget()
    # Enable main buttons
    start_button.configure(state="normal")
    stop_button.configure(state="disabled")
    pause_button.configure(state="disabled", text="Pause Recording")  # --- New Feature: Disable Pause Button ---
    next_button.configure(state="normal")
    prev_button.configure(state="normal")
    save_button.configure(state="normal")
    skip_button.configure(state="normal")
    update_add_more_button_state()  # Update 'Add More' button state

# Function when user accepts the response
def accept_response():
    hide_decision_frame()
    next_question_func()

# Function when user wants to redo the response
def redo_response():
    hide_decision_frame()
    # Clear the response for the current question
    current_q = questions[current_question]
    transcribed_responses[current_q] = ""
    update_resume_preview()
    update_add_more_button_state()
    start_recording_thread(add_more=False)

# Function when user wants to add more to the response
def add_more_response():
    hide_decision_frame()
    start_recording_thread(add_more=True)
    
    
def recover_last_attempt():
    global start_frame
    # Define the autosave file path within data_directory
    autosave_file_path = get_autosave_file_path()

    # Check if the autosave file exists
    if os.path.exists(autosave_file_path):
        loaded = load_autosave_data(autosave_file_path)  # Uses the updated path
        if loaded:
            if current_mode == 'manual':
                if start_frame is not None:
                    start_frame.destroy()
                    start_frame = None
                if 'question_label' in globals():
                    update_question()
                    update_resume_preview()
                else:
                    initialize_manual_gui()
                    update_question()
                    update_resume_preview()
            elif current_mode == 'voice':
                if start_frame is not None:
                    start_frame.destroy()
                    start_frame = None
                if 'question_label' in globals():
                    update_question_voice_mode(current_question)
                    update_resume_preview()
                else:
                    initialize_voice_gui()
                    if voice_thread is None or not voice_thread.is_alive():
                        voice_thread = threading.Thread(target=voice_walkthrough, daemon=True)
                        voice_thread.start()
                update_resume_preview()
            messagebox.showinfo("Recovered", "Your last attempt has been recovered.")
        else:
            messagebox.showerror("Error", "Failed to load the last attempt.")
    else:
        messagebox.showinfo("No Data", "No autosave data found to recover.")




# ========================== Settings Implementation ==========================

# Global dictionary to store TTS settings and language selection
tts_settings = {
    'rate': 200,  # Default speech rate
    'voice': None,  # To be set after initializing engine
    'language': 'English'  # Default language
}

# Mapping of languages to their codes
language_mapping = {
    'English': 'en',
    'Spanish': 'es',
    'French': 'fr'
}

selected_language = tts_settings['language']
selected_language_code = language_mapping[selected_language]

# Initialize the pyttsx3 engine globally
engine = pyttsx3.init()
# Set default properties
engine.setProperty('rate', tts_settings['rate'])

# Get available voices and set default
voices = engine.getProperty('voices')
if voices:
    # Default to first voice that matches the selected language, else first voice
    selected_voice = next((voice for voice in voices if selected_language_code in voice.languages or selected_language_code in voice.id.lower()), voices[0])
    tts_settings['voice'] = selected_voice
    engine.setProperty('voice', selected_voice.id)
else:
    messagebox.showerror("Error", "No voices available for the text-to-speech engine.")

# Function to open the Settings window
def open_settings():
    settings_window = ctk.CTkToplevel(root)
    settings_window.title("Settings")
    settings_window.geometry("400x400")
    settings_window.resizable(False, False)
    
    # Set the settings window to be always on top and bring it to focus
    settings_window.attributes("-topmost", True)
    settings_window.focus_force()
    
    # --- Microphone Selection ---
    def get_microphone_list():
        p = pyaudio.PyAudio()
        mic_list = []

        for i in range(p.get_device_count()):
            device_info = p.get_device_info_by_index(i)
            # Only include devices with input channels that are marked as available
            if device_info['maxInputChannels'] > 0 and device_info['hostApi'] == 0:  # 'hostApi' 0 is usually the primary Windows audio API
                mic_list.append((i, device_info['name']))

        p.terminate()
        return mic_list
    
    # Label for Microphone Selection
    mic_label = ctk.CTkLabel(settings_window, text="Select Microphone:", font=("Helvetica", 14))
    mic_label.pack(pady=(20, 5))

    # Get the list of microphones
    mic_list = get_microphone_list()
    mic_names = [name for index, name in mic_list]

    # Variable to hold the selected microphone
    selected_mic_var = ctk.StringVar(value=mic_names[0] if mic_names else "No Microphone Available")

    def update_microphone(mic_name):
        global selected_microphone_index
        for index, name in mic_list:
            if name == mic_name:
                selected_microphone_index = index
                break
        else:
            selected_microphone_index = None  # No microphone selected

    mic_dropdown = ctk.CTkOptionMenu(settings_window, values=mic_names, variable=selected_mic_var, command=update_microphone)
    mic_dropdown.pack(pady=5, padx=20, fill="x")

    # Set the initial selected microphone index
    update_microphone(selected_mic_var.get())

    # Label for Speech Rate
    rate_label = ctk.CTkLabel(settings_window, text="Speech Rate:", font=("Helvetica", 14))
    rate_label.pack(pady=(20, 5))

    # Slider for Speech Rate
    rate_slider = ctk.CTkSlider(settings_window, from_=100, to=300, number_of_steps=200, command=lambda val: update_rate(val))
    rate_slider.set(tts_settings['rate'])
    rate_slider.pack(pady=5, padx=20, fill="x")

    # Display current rate value
    rate_value_label = ctk.CTkLabel(settings_window, text=f"{tts_settings['rate']} words per minute", font=("Helvetica", 12))
    rate_value_label.pack(pady=5)

    # Update label when slider moves
    def on_rate_change(val):
        rate = int(float(val))
        rate_value_label.configure(text=f"{rate} words per minute")

    rate_slider.configure(command=lambda val: [update_rate(val), on_rate_change(val)])

    # Label for Voice Selection
    voice_label = ctk.CTkLabel(settings_window, text="Select Voice:", font=("Helvetica", 14))
    voice_label.pack(pady=(20, 5))

    # Dropdown for Voice Selection
    voice_options = [voice.name for voice in voices]
    selected_voice_name = ctk.StringVar(value=tts_settings['voice'].name if tts_settings['voice'] else "")

    voice_dropdown = ctk.CTkOptionMenu(settings_window, values=voice_options, variable=selected_voice_name, command=lambda val: update_voice(val))
    voice_dropdown.pack(pady=5, padx=20, fill="x")

    # Label for Language Selection
    language_label = ctk.CTkLabel(settings_window, text="Select Language:", font=("Helvetica", 14))
    language_label.pack(pady=(20, 5))

    # Dropdown for Language Selection
    language_options = list(language_mapping.keys())
    selected_language_var = ctk.StringVar(value=tts_settings['language'])

    language_dropdown = ctk.CTkOptionMenu(settings_window, values=language_options, variable=selected_language_var, command=lambda val: update_language(val))
    language_dropdown.pack(pady=5, padx=20, fill="x")

    # Close Button
    close_button = ctk.CTkButton(settings_window, text="Close", command=settings_window.destroy)
    close_button.pack(pady=20)

# Function to update speech rate
def update_rate(val):
    rate = int(float(val))
    tts_settings['rate'] = rate
    engine.setProperty('rate', rate)

# Function to update voice
def update_voice(voice_name):
    selected = next((voice for voice in voices if voice.name == voice_name), None)
    if selected:
        tts_settings['voice'] = selected
        engine.setProperty('voice', selected.id)
    else:
        messagebox.showerror("Error", "Selected voice not found.")

# Function to update language
def update_language(language):
    global selected_language, selected_language_code
    if language in language_mapping:
        selected_language = language
        selected_language_code = language_mapping[language]
        tts_settings['language'] = language
        # Optionally, you can also set the voice to match the language
        # This requires that voices are tagged with their languages
        # For simplicity, we're not changing the voice based on language
    else:
        messagebox.showerror("Error", "Selected language not supported.")
        
        
def get_current_date_string():
    return datetime.now().strftime("%m%d%Y")

def get_unique_filename(base_name, extension, directory=data_directory):
    date_str = get_current_date_string()
    filename = f"{base_name}-{date_str}.{extension}"
    counter = 1
    while os.path.exists(os.path.join(directory, filename)):
        filename = f"{base_name}-{date_str}_{counter}.{extension}"
        counter += 1
    return filename


# ========================== End of Settings Implementation ==========================

# GUI Setup using customtkinter
ctk.set_appearance_mode("System")  # Options: "System" (default), "Dark", "Light"
ctk.set_default_color_theme("dark-blue")  # Options: "blue" (default), "green", "dark-blue"

root = ctk.CTk()
root.title("Resume Generator")
root.geometry("1500x900")
root.resizable(True, True)  # Make the window resizable

# ========================== Menu Bar Implementation ==========================

# Help Button
def create_menu_bar():
    menu_bar = ctk.CTkFrame(root, corner_radius=0)
    menu_bar.pack(side="top", fill="x")


    # Set Save Directory Button
    def set_save_directory():
        global save_directory, data_directory
        selected_dir = filedialog.askdirectory(initialdir=save_directory)
        if selected_dir:
            save_directory = selected_dir
            data_directory = os.path.join(save_directory, 'data')
            os.makedirs(data_directory, exist_ok=True)  # Ensure the data subdirectory exists
            messagebox.showinfo("Save Directory Set", f"Files will now be saved in:\n{save_directory}")
        else:
            # User cancelled the dialog, keep the existing directory
            pass


    save_dir_button = ctk.CTkButton(
        menu_bar,
        text="Set Save Directory",
        command=set_save_directory,
        width=150,
        height=30,
        fg_color="transparent",
        hover_color="gray70",
        text_color="white",
        font=("Helvetica", 12)
    )
    save_dir_button.pack(side="left", padx=10, pady=5)
    
    
    # Settings Button
    settings_button = ctk.CTkButton(
        menu_bar,
        text="Settings",
        command=open_settings,
        width=100,
        height=30,
        fg_color="transparent",
        hover_color="gray70",
        text_color="white",
        font=("Helvetica", 12)
    )
    settings_button.pack(side="left", padx=10, pady=5)
    
    # Recover Last Attempt Button
    recover_button = ctk.CTkButton(
        menu_bar,
        text="Recover Last Attempt",
        command=recover_last_attempt,
        width=150,
        height=30,
        fg_color="transparent",
        hover_color="gray70",
        text_color="white",
        font=("Helvetica", 12)
    )
    recover_button.pack(side="left", padx=10, pady=5)

    # Help Button
    def show_help():
        # Create a new top-level window for the Help Menu
        help_window = ctk.CTkToplevel(root)
        help_window.title("Help")
        help_window.geometry("800x600")
        help_window.resizable(False, False)  # Fixed size for consistency
        
    
        # Set the settings window to be always on top and bring it to focus
        help_window.attributes("-topmost", True)
        help_window.focus_force()

        # Create a frame to hold all help content with padding
        help_frame = ctk.CTkFrame(help_window, corner_radius=10)
        help_frame.pack(pady=20, padx=20, fill="both", expand=True)

        # Add a title label
        title_label = ctk.CTkLabel(
            help_frame,
            text="Resume Generator Help",
            font=("Helvetica", 20, "bold"),
            anchor="w"
        )
        title_label.pack(pady=(0, 10), anchor="w")

        # Create a scrollable frame in case content overflows
        scroll_frame = ctk.CTkScrollableFrame(help_frame, width=760, height=500)
        scroll_frame.pack(pady=10, padx=5, fill="both", expand=True)

        # Helper function to add section titles
        def add_section_title(text):
            section_label = ctk.CTkLabel(
                scroll_frame,
                text=text,
                font=("Helvetica", 16, "bold"),
                anchor="w"
            )
            section_label.pack(pady=(10, 5), anchor="w")

        # Helper function to add bullet points
        def add_bullet_point(text):
            bullet_label = ctk.CTkLabel(
                scroll_frame,
                text=f"• {text}",
                font=("Helvetica", 12),
                anchor="w"
            )
            bullet_label.pack(pady=2, anchor="w")

        # Helper function to add numbered lists
        def add_numbered_list(text, number):
            numbered_label = ctk.CTkLabel(
                scroll_frame,
                text=f"{number}. {text}",
                font=("Helvetica", 12),
                anchor="w"
            )
            numbered_label.pack(pady=2, anchor="w")

        # -------------------- Help Content Sections --------------------

        # How to Use Section
        add_section_title("How to Use")
        add_numbered_list("Manual Walkthrough: Enter your details manually.", 1)
        add_numbered_list("Voice Walkthrough: Answer questions verbally using your microphone.", 2)

        # Keyboard Shortcuts Section
        add_section_title("Keyboard Shortcuts")
        add_bullet_point("Ctrl+R: Start/Stop voice recording")
        add_bullet_point("Ctrl+N: Next question")
        add_bullet_point("Ctrl+P: Previous question")
        add_bullet_point("Ctrl+S: Save your resume")
        add_bullet_point("Ctrl+A: Add more details")
        add_bullet_point("Ctrl+Q: Quit the app")

        # Settings Section
        add_section_title("Settings")
        add_bullet_point("Adjust speech speed and choose a Text-to-Speech (TTS) voice.")
        add_bullet_point("Select a transcription language for better accuracy.")
        
        # File Storage Information Section
        add_section_title("File Storage Information")
        add_bullet_point("Default Save Location: All files are saved in your Downloads folder in a folder named 'VoiceResumeMaker'.")
        add_bullet_point("Data Folder: Intermediary files and responses are saved in a subfolder named 'data' inside 'VoiceResumeMaker'.")
        add_bullet_point("Final Files: The completed PDF and Word documents are saved directly into the 'VoiceResumeMaker' folder for easy access.")

        # Microphone Setup Section
        add_section_title("Microphone Setup")
        add_bullet_point("Ensure your microphone is properly connected for the voice mode to function correctly.")
        add_bullet_point("Go into settings and select the current microphone that is connected, if there is none it will be defaulted to the system microphone.")

        # Recording Guidelines Section
        add_section_title("Recording Guidelines")
        add_bullet_point("Maximum File Size: 25 MB per recording.")
        add_bullet_point("Estimated Duration: Approximately 30 seconds per 1 MB, allowing up to 12.5 minutes per recording.")
        add_bullet_point("Monitoring Size: The application displays the recording size in MB before sending it for transcription.")

        # Recording Tips Section
        add_section_title("Recording Tips")
        add_bullet_point("Speak Clearly: Enunciate your words for better transcription accuracy.")
        add_bullet_point("Minimize Background Noise: Reduce ambient sounds to avoid interference.")
        add_bullet_point("Stay Within Limits: Keep recordings under the 12-minute threshold to prevent errors.")

        # Contact Information Section
        add_section_title("Need Help?")
        contact_label = ctk.CTkLabel(
            scroll_frame,
            text="For assistance or questions, please contact our support team at support@example.com.",
            font=("Helvetica", 12),
            anchor="w",
            wraplength=700
        )
        contact_label.pack(pady=(2, 10), anchor="w")

        # Optional: Add a close button at the bottom
        close_button = ctk.CTkButton(
            help_frame,
            text="Close",
            command=help_window.destroy,
            width=100,
            height=40,
            fg_color="blue",
            text_color="white",
            font=("Helvetica", 12)
        )
        close_button.pack(pady=10)

    help_button = ctk.CTkButton(
        menu_bar,
        text="Help",
        command=show_help,
        width=100,
        height=30,
        fg_color="transparent",
        hover_color="gray70",
        text_color="white",
        font=("Helvetica", 12)
    )
    help_button.pack(side="left", padx=10, pady=5)

    # ========================== Keyboard Shortcuts Implementation ==========================

    def on_shortcut(event):
        # Mapping keyboard shortcuts to functions
        # Ctrl+R: Start/Stop Recording
        # Ctrl+N: Next Question
        # Ctrl+P: Previous Question
        # Ctrl+S: Save Resume
        # Ctrl+A: Add More
        # Ctrl+Q: Quit Application
        if event.state & 0x0004:  # Ctrl key
            if event.keysym.lower() == 'r':
                # Toggle Recording
                if stop_button.cget('state') == 'normal':
                    stop_recording()
                else:
                    start_recording_thread(add_more=False)
            elif event.keysym.lower() == 'n':
                next_question_func()
            elif event.keysym.lower() == 'p':
                prev_question_func()
            elif event.keysym.lower() == 's':
                save_resume()
            elif event.keysym.lower() == 'a':
                add_more_response()
            elif event.keysym.lower() == 'q':
                root.quit()

    root.bind_all("<Control-r>", on_shortcut)  # Ctrl+R
    root.bind_all("<Control-n>", on_shortcut)  # Ctrl+N
    root.bind_all("<Control-p>", on_shortcut)  # Ctrl+P
    root.bind_all("<Control-s>", on_shortcut)  # Ctrl+S
    root.bind_all("<Control-a>", on_shortcut)  # Ctrl+A
    root.bind_all("<Control-q>", on_shortcut)  # Ctrl+Q

    # ========================== End of Keyboard Shortcuts Implementation ==========================

create_menu_bar()

# Load icons (Ensure the icons are in the same directory or provide the correct path)
try:
    manual_icon = ctk.CTkImage(Image.open(resource_path("icons/manual_icon.png")), size=(20, 20))
    voice_icon = ctk.CTkImage(Image.open(resource_path("icons/voice_icon.png")), size=(20, 20))
    start_icon = ctk.CTkImage(Image.open(resource_path("icons/start_icon.png")), size=(20, 20))
    stop_icon = ctk.CTkImage(Image.open(resource_path("icons/stop_icon.png")), size=(20, 20))
    next_icon = ctk.CTkImage(Image.open(resource_path("icons/next_icon.png")), size=(20, 20))
    prev_icon = ctk.CTkImage(Image.open(resource_path("icons/prev_icon.png")), size=(20, 20))
    save_icon = ctk.CTkImage(Image.open(resource_path("icons/save_icon.png")), size=(20, 20))
    skip_icon = ctk.CTkImage(Image.open(resource_path("icons/skip_icon.png")), size=(20, 20))
    accept_icon = ctk.CTkImage(Image.open(resource_path("icons/accept_icon.png")), size=(20, 20))
    redo_icon = ctk.CTkImage(Image.open(resource_path("icons/redo_icon.png")), size=(20, 20))
    add_more_icon = ctk.CTkImage(Image.open(resource_path("icons/add_more_icon.png")), size=(20, 20))
    pause_icon = ctk.CTkImage(Image.open(resource_path("icons/pause_icon.png")), size=(20, 20))
except Exception as e:
    print(f"Icon loading failed: {e}")
    manual_icon = None
    voice_icon = None
    start_icon = None
    stop_icon = None
    next_icon = None
    prev_icon = None
    save_icon = None
    skip_icon = None
    accept_icon = None
    redo_icon = None
    add_more_icon = None
    pause_icon = None

# Progress Bar
progress_bar = ctk.CTkProgressBar(root, width=400, height=20)
progress_bar.set(0)  # Initial progress
progress_bar.pack(pady=(20, 10), padx=20)


# Function to show the start page
def show_start_page():
    global start_frame  
    start_frame = ctk.CTkFrame(root, corner_radius=10)
    start_frame.pack(pady=10, padx=20, fill="both", expand=True)

    # Welcome Image or Logo (Ensure the image exists)
    try:
        logo_image = ctk.CTkImage(Image.open(resource_path("icons/logo.png")), size=(100, 100))
        logo_label = ctk.CTkLabel(start_frame, image=logo_image, text="")
        logo_label.image = logo_image  # Keep a reference
        logo_label.pack(pady=(20, 10))
    except Exception as e:
        print(f"Logo loading failed: {e}")

    welcome_label = ctk.CTkLabel(start_frame, text="Welcome to the Resume Generator", font=("Helvetica", 26, "bold"))
    welcome_label.pack(pady=10)

    instruction_label = ctk.CTkLabel(start_frame, text="Please select an option to proceed", font=("Helvetica", 16))
    instruction_label.pack(pady=10)

    # Manual Walkthrough Button
    manual_button = ctk.CTkButton(
        start_frame,
        text="Manual Walkthrough",
        image=manual_icon,
        compound="left",
        command=lambda: start_manual_mode(start_frame),
        width=200,
        height=50,
        font=("Helvetica", 14)
    )
    manual_button.pack(pady=10)

    # Voice Assisted Walkthrough Button
    voice_button = ctk.CTkButton(
        start_frame,
        text="Voice Assisted Walkthrough",
        image=voice_icon,
        compound="left",
        command=lambda: start_voice_mode(start_frame),
        width=200,
        height=50,
        font=("Helvetica", 14)
    )
    voice_button.pack(pady=10)

# Function to load autosave data
def load_autosave_data(autosave_file_path):
    global transcribed_responses, current_question, current_mode
    try:
        with open(autosave_file_path, 'r') as f:
            data = json.load(f)
            transcribed_responses = data.get('transcribed_responses', {})
            current_question = data.get('current_question', 0)
            current_mode = data.get('current_mode', 'manual')
            return True
    except Exception as e:
        messagebox.showerror("Error", f"Failed to load autosave data: {e}")
        return False

        
# Function to autosave data periodically
def autosave_data():
    data = {
        'transcribed_responses': transcribed_responses,
        'current_question': current_question,
        'current_mode': current_mode
    }
    autosave_file = get_autosave_file_path()
    try:
        with open(autosave_file, 'w') as f:
            json.dump(data, f)
    except Exception as e:
        print(f"Failed to autosave data: {e}")
    # Schedule the next autosave
    root.after(30000, autosave_data)  # Autosave every 30 seconds




# Function to start manual mode with voice recording
def start_manual_mode(start_frame):
    global current_mode
    current_mode = 'manual'
    start_frame.destroy()
    start_frame = None  # Add this line
    initialize_manual_gui()
    # Start autosave
    autosave_data()

# Function to start voice mode
def start_voice_mode(start_frame):
    global voice_thread, current_mode
    current_mode = 'voice'
    start_frame.destroy()
    start_frame = None  # Add this line
    initialize_voice_gui()
    if voice_thread is None or not voice_thread.is_alive():
        voice_thread = threading.Thread(target=voice_walkthrough, daemon=True)
        voice_thread.start()
    # Start autosave
    autosave_data()


# Function to initialize the voice GUI
def initialize_voice_gui():
    global status_label, question_label, transcribed_label, resume_preview_textbox

    # Create a frame for the voice mode
    voice_frame = ctk.CTkFrame(root, corner_radius=10)
    voice_frame.pack(pady=10, padx=20, fill="both", expand=True)

    # Frame for Status Label
    status_frame = ctk.CTkFrame(voice_frame, corner_radius=10)
    status_frame.pack(pady=10, padx=20, fill="x")

    # Status Label (Centered and with increased wraplength)
    status_label = ctk.CTkLabel(
        status_frame,
        text="Status: Starting Voice Walkthrough...",
        wraplength=900,  # Increased for better readability
        font=("Helvetica", 14, "italic"),
        anchor="center"  # Center-align the text
    )
    status_label.pack(pady=5, padx=20, anchor="center")  # Centered within the frame

    # Frame for Question Label
    question_frame = ctk.CTkFrame(voice_frame, corner_radius=10)
    question_frame.pack(pady=10, padx=20, fill="x")

    # Label to display the current question (Centered and with larger font)
    question_label = ctk.CTkLabel(
        question_frame,
        text="",
        wraplength=900,  # Increased wraplength
        font=("Helvetica", 20, "bold"),  # Larger font and bold
        anchor="center"  # Center-align the text
    )
    question_label.pack(pady=(20, 10), padx=20, fill="x")

    # Frame for Transcription Label
    transcription_frame = ctk.CTkFrame(voice_frame, corner_radius=10)
    transcription_frame.pack(pady=10, padx=20, fill="x")

    # Label to display transcription in a distinct block (Centered and with padding)
    transcribed_label = ctk.CTkLabel(
        transcription_frame,
        text="Transcription: ",
        wraplength=800,
        font=("Helvetica", 16),  # Larger font for readability
        anchor="center",  # Center-align the transcription
        padx=20,  # Add padding for a spacious look
        pady=10
    )
    transcribed_label.pack(fill="x")

    # Resume Preview Frame
    preview_frame = ctk.CTkFrame(voice_frame, corner_radius=10)
    preview_frame.pack(pady=10, padx=20, fill="both", expand=True)

    # Resume Preview Label
    resume_preview_label = ctk.CTkLabel(
        preview_frame,
        text="Resume Preview:",
        font=("Helvetica", 16, "bold")
    )
    resume_preview_label.pack(pady=(10, 5), padx=10, anchor="w")

    # Resume Preview Textbox
    resume_preview_textbox = ctk.CTkTextbox(
        preview_frame,
        wrap="word",
        font=("Consolas", 12)
    )
    resume_preview_textbox.pack(pady=5, padx=10, fill="both", expand=True)
    resume_preview_textbox.configure(state="disabled")  # Make it read-only

    # Update the GUI with the current question
    update_question_voice_mode(current_question)


def update_add_more_button_state():
    question = questions[current_question]
    if question in transcribed_responses and transcribed_responses[question].strip():
        # There is an existing response, enable 'Add More' button
        add_more_main_button.configure(state="normal")
    else:
        # No response yet, disable 'Add More' button
        add_more_main_button.configure(state="disabled")


# Function to initialize the manual input GUI with voice recording
def initialize_manual_gui():
    global question_label, transcribed_label
    global start_button, stop_button, pause_button, next_button, prev_button, save_button, skip_button, generate_button
    global status_label, resume_preview_textbox, decision_frame
    global add_more_main_button
    global timer_label

    # Main frame for manual mode
    manual_frame = ctk.CTkFrame(root, corner_radius=10)
    manual_frame.pack(pady=10, padx=20, fill="both", expand=True)

    # Frame for question and transcription
    question_frame = ctk.CTkFrame(manual_frame, corner_radius=10)
    question_frame.pack(pady=10, padx=20, fill="x")

    # Label to display the current question
    question_label = ctk.CTkLabel(
        question_frame,
        text="",
        wraplength=900,  # Increase wrap length to span more horizontally
        font=("Helvetica", 20, "bold"),  # Larger font and bold for readability
        anchor="center"  # Center-align the text
    )
    question_label.pack(pady=(20, 10), padx=20, fill="x")
    
    # Dedicated Frame for Transcription Block
    transcription_frame = ctk.CTkFrame(manual_frame, corner_radius=10)
    transcription_frame.pack(pady=10, padx=20, fill="x")

    # Label to display transcription in a distinct block
    transcribed_label = ctk.CTkLabel(
        transcription_frame,
        text="Transcription: ",
        wraplength=800,
        font=("Helvetica", 16),  # Larger font for readability
        anchor="center",  # Center-align the transcription
        padx=20,  # Add padding for a spacious look
        pady=10
    )
    transcribed_label.pack(fill="x")

    # Frame for recording controls
    recording_frame = ctk.CTkFrame(manual_frame, corner_radius=10)
    recording_frame.pack(pady=10, padx=20, fill="x")
    
    # Configure columns to expand equally in the recording_frame
    for i in range(5):
        recording_frame.columnconfigure(i, weight=1)

    # Timer Label
    timer_label = ctk.CTkLabel(
        recording_frame,
        text="Recording Time: 00:00:00",
        font=("Helvetica", 12)
    )
    timer_label.grid(row=0, column=0, padx=5, pady=10, sticky="ew")

    # Start Recording Button
    start_button = ctk.CTkButton(
        recording_frame,
        text="Start Recording",
        image=start_icon,
        compound="left",
        command=lambda: start_recording_thread(add_more=False),
        width=140,
        height=40,
        font=("Helvetica", 12)
    )
    start_button.grid(row=0, column=1, padx=5, pady=10, sticky="ew")

    # Pause Recording Button
    pause_button = ctk.CTkButton(
        recording_frame,
        text="Pause Recording",
        image=pause_icon,
        compound="left",
        command=toggle_pause_recording,
        width=140,
        height=40,
        font=("Helvetica", 12),
        state="disabled",
        fg_color="green"
    )
    pause_button.grid(row=0, column=2, padx=5, pady=10, sticky="ew")

    # Stop Recording Button
    stop_button = ctk.CTkButton(
        recording_frame,
        text="Stop Recording",
        image=stop_icon,
        compound="left",
        command=stop_recording,
        width=140,
        height=40,
        font=("Helvetica", 12),
        state="disabled"
    )
    stop_button.grid(row=0, column=3, padx=5, pady=10, sticky="ew")

    # Add More Button
    add_more_main_button = ctk.CTkButton(
        recording_frame,
        text="Add More",
        image=add_more_icon,
        compound="left",
        command=add_more_response,
        width=140,
        height=40,
        font=("Helvetica", 12)
    )
    add_more_main_button.grid(row=0, column=4, padx=5, pady=10, sticky="ew")

    # Frame for navigation controls
    navigation_frame = ctk.CTkFrame(manual_frame, corner_radius=10)
    navigation_frame.pack(pady=10, padx=20, fill="x")
    
    # Configure columns to expand equally in the navigation_frame
    for i in range(5):
        navigation_frame.columnconfigure(i, weight=1)

    # Previous Button
    prev_button = ctk.CTkButton(
        navigation_frame,
        text="Previous",
        image=prev_icon,
        compound="left",
        command=prev_question_func,
        width=120,
        height=40,
        font=("Helvetica", 12)
    )
    prev_button.grid(row=0, column=0, padx=5, pady=10, sticky="ew")

    # Next Button
    next_button = ctk.CTkButton(
        navigation_frame,
        text="Next",
        image=next_icon,
        compound="left",
        command=next_question_func,
        width=120,
        height=40,
        font=("Helvetica", 12)
    )
    next_button.grid(row=0, column=1, padx=5, pady=10, sticky="ew")

    # Skip Button
    skip_button = ctk.CTkButton(
        navigation_frame,
        text="Skip",
        image=skip_icon,
        compound="left",
        command=skip_question_func,
        width=120,
        height=40,
        font=("Helvetica", 12)
    )
    skip_button.grid(row=0, column=2, padx=5, pady=10, sticky="ew")

    # Save Resume Button
    save_button = ctk.CTkButton(
        navigation_frame,
        text="Save Resume",
        image=save_icon,
        compound="left",
        command=save_resume,
        width=150,
        height=40,
        font=("Helvetica", 12)
    )
    save_button.grid(row=0, column=0, padx=10, pady=10)
    save_button.grid_remove()  # Hides the button after adding it to the layout

    # Generate Formatted Resume Button
    generate_button = ctk.CTkButton(
        navigation_frame,
        text="Generate Formatted Resume",
        image=save_icon,
        compound="left",
        command=generate_and_save_formatted_resume,
        width=220,
        height=40,
        font=("Helvetica", 12)
    )
    generate_button.grid(row=0, column=1, padx=10, pady=10)
    generate_button.grid_remove() # Hides the button after adding it to the layout

    # Create Finished Resume Button
    create_resume_button = ctk.CTkButton(
        navigation_frame,
        text="Create Finished Resume",
        image=save_icon,
        compound="left",
        command=create_finished_resume,
        width=200,
        height=40,
        font=("Helvetica", 12)
    )
    create_resume_button.grid(row=0, column=3, padx=5, pady=10, sticky="ew")

    # Clear Data Button
    clear_data_button = ctk.CTkButton(
        navigation_frame,
        text="Clear Data",
        image=None,
        compound="left",
        command=clear_data,
        width=150,
        height=40,
        font=("Helvetica", 12),
        fg_color="red",
        hover_color="darkred"
    )
    clear_data_button.grid(row=0, column=4, padx=5, pady=10, sticky="ew")

    # Status Label
    status_label = ctk.CTkLabel(
        manual_frame,
        text="Status: Ready",
        wraplength=760,
        font=("Helvetica", 14, "italic")
    )
    status_label.pack(pady=10, padx=20, anchor="w")

    # Decision Frame (initially hidden)
    decision_frame = ctk.CTkFrame(manual_frame, corner_radius=10)
    # Decision Frame Buttons
    accept_button = ctk.CTkButton(
        decision_frame,
        text="Accept",
        image=accept_icon,
        compound="left",
        command=accept_response,
        width=120,
        height=40,
        font=("Helvetica", 12)
    )
    accept_button.pack(side="left", padx=10, pady=10)

    redo_button = ctk.CTkButton(
        decision_frame,
        text="Redo",
        image=redo_icon,
        compound="left",
        command=redo_response,
        width=120,
        height=40,
        font=("Helvetica", 12)
    )
    redo_button.pack(side="left", padx=10, pady=10)

    add_more_button = ctk.CTkButton(
        decision_frame,
        text="Add More",
        image=add_more_icon,
        compound="left",
        command=add_more_response,
        width=120,
        height=40,
        font=("Helvetica", 12)
    )
    add_more_button.pack(side="left", padx=10, pady=10)

    # Resume Preview Frame
    preview_frame = ctk.CTkFrame(manual_frame, corner_radius=10)
    preview_frame.pack(pady=10, padx=20, fill="both", expand=True)

    # Resume Preview Label
    resume_preview_label = ctk.CTkLabel(
        preview_frame,
        text="Resume Preview:",
        font=("Helvetica", 16, "bold")
    )
    resume_preview_label.pack(pady=(10, 5), padx=10, anchor="w")

    # Resume Preview Textbox
    resume_preview_textbox = ctk.CTkTextbox(
        preview_frame,
        wrap="word",
        font=("Consolas", 20)
    )
    resume_preview_textbox.pack(pady=5, padx=10, fill="both", expand=True)
    resume_preview_textbox.configure(state="disabled")  # Make it read-only

    # Initialize the first question
    update_question()
    update_resume_preview()
    # Start autosave
    autosave_data()

# Start the application
show_start_page()

# Start the GUI event loop
root.mainloop()