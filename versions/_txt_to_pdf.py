# TODO :
#  - After the first walkthrough is complleted save the file and use that for creating the resume itself, then modify that file and see how the template reacts
#  - Make it so that it firsts get placed into a word doc (formatted), and is then converted to a pdf for the final result
#  - allow the word doc to be downloaded as well so personal modification can be done as well 
#  - check if it is possible to get a nicer voice to talk to you in the voice assisted version


import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert
import os

def parse_resume_text(text_file):
    """
    Parses the resume text file and structures the data into a dictionary.
    """
    with open(text_file, 'r', encoding='utf-8') as file:
        lines = [line.strip() for line in file.readlines()]

    data = {
        "name": "",
        "phone": "",
        "email": "",
        "address": "",
        "professional_summary": "",
        "skills": [],
        "certifications": [],
        "achievements": [],
        "work_experience": [],
        "education": [],
        "interests": [],
        "extracurricular_activities": [],
        "volunteer_experience": [],
        "professional_associations": [],
        "references": ""
    }

    sections = [
        "Professional Summary", "Skills", "Certifications and Training",
        "Professional Achievements", "Work Experience", "Education",
        "Interests", "Extracurricular Activities", "Volunteer Experience",
        "Professional Associations", "References"
    ]
    
    # Define possible position labels
    position_labels = ["Position:", "Standing:", "Role:", "Title:"]

    current_section = None
    i = 0
    n = len(lines)

    print("Starting to parse the resume text file...")

    # Parse header (Name, Phone, Email, Address)
    if i < n:
        data["name"] = lines[i]
        print(f"Name: {data['name']}")
        i += 1
    if i < n:
        data["phone"] = lines[i]
        print(f"Phone: {data['phone']}")
        i += 1
    if i < n:
        data["email"] = lines[i]
        print(f"Email: {data['email']}")
        i += 1
    if i < n:
        data["address"] = lines[i]
        print(f"Address: {data['address']}")
        i += 1

    # Parse sections
    while i < n:
        line = lines[i]
        print(f"Processing line {i}: '{line}' (length: {len(line)})")

        if line in sections:
            current_section = line
            print(f"--- Entered section: {current_section} ---")
            i += 1
            continue

        if not current_section:
            print(f"No current section. Skipping line {i}.")
            i += 1
            continue

        try:
            if current_section == "Professional Summary":
                summary = []
                while i < n and lines[i] not in sections:
                    if lines[i]:  # Skip empty lines
                        summary.append(lines[i])
                        print(f"Appending to Professional Summary: {lines[i]}")
                    i += 1
                data["professional_summary"] = ' '.join(summary)
                print(f"Professional Summary: {data['professional_summary']}")
                # Reset section
                current_section = None
                print(f"--- Exiting section: Professional Summary ---")

            elif current_section == "Skills":
                while i < n and lines[i].startswith("- "):
                    skill = lines[i][2:].strip()
                    data["skills"].append(skill)
                    print(f"Appending Skill: {skill}")
                    i += 1
                # Reset section
                current_section = None
                print(f"--- Exiting section: Skills ---")

            elif current_section == "Certifications and Training":
                while i < n and lines[i].startswith("- "):
                    cert = lines[i][2:].strip()
                    data["certifications"].append(cert)
                    print(f"Appending Certification: {cert}")
                    i += 1
                # Reset section
                current_section = None
                print(f"--- Exiting section: Certifications and Training ---")

            elif current_section == "Professional Achievements":
                while i < n and lines[i].startswith("- "):
                    achievement = lines[i][2:].strip()
                    data["achievements"].append(achievement)
                    print(f"Appending Achievement: {achievement}")
                    i += 1
                # Reset section
                current_section = None
                print(f"--- Exiting section: Professional Achievements ---")

            elif current_section == "Work Experience":
                while i < n and lines[i] not in sections:
                    if not lines[i]:  # Skip empty lines
                        print(f"Encountered empty line in Work Experience. Moving to next line.")
                        i += 1
                        continue

                    # Read company name
                    company = lines[i]
                    print(f"Company: {company}")
                    i += 1

                    # Initialize position and description
                    position = ""
                    description = []

                    # Check for position labels
                    for label in position_labels:
                        if i < n and lines[i].startswith(label):
                            position = lines[i].replace(label, "").strip()
                            print(f"Position: {position}")
                            i += 1
                            break  # Stop checking after finding the first matching label

                    # Collect job description
                    while i < n and lines[i] not in sections:
                        current_line = lines[i]

                        # Determine if the current line indicates the start of a new work experience
                        is_new_company = (
                            not any(current_line.startswith(l) for l in position_labels) and
                            not current_line.startswith("- ") and
                            not current_line.lower().startswith("as a") and
                            not current_line.lower().startswith("as an")
                        )

                        if is_new_company:
                            print(f"Detected new company or invalid line for description: '{current_line}'. Ending current job description.")
                            break  # New company detected, stop description collection

                        # Otherwise, it's part of the current job's description
                        description.append(current_line)
                        print(f"Appending to Job Description: {current_line}")
                        i += 1

                    # Append the job to work_experience
                    job = {
                        "company": company,
                        "position": position,
                        "description": ' '.join(description)
                    }
                    data["work_experience"].append(job)
                    print(f"Added Work Experience: {job}")

                # Reset section
                current_section = None
                print(f"--- Exiting section: Work Experience ---")

            elif current_section == "Education":
                while i < n and lines[i] not in sections:
                    if lines[i]:  # Skip empty lines
                        edu = lines[i]
                        data["education"].append(edu)
                        print(f"Appending Education: {edu}")
                    i += 1
                # Reset section
                current_section = None
                print(f"--- Exiting section: Education ---")

            elif current_section == "Interests":
                while i < n and lines[i].startswith("- "):
                    interest = lines[i][2:].strip()
                    data["interests"].append(interest)
                    print(f"Appending Interest: {interest}")
                    i += 1
                # Reset section
                current_section = None
                print(f"--- Exiting section: Interests ---")

            elif current_section == "Extracurricular Activities":
                while i < n and lines[i].startswith("- "):
                    activity = lines[i][2:].strip()
                    data["extracurricular_activities"].append(activity)
                    print(f"Appending Extracurricular Activity: {activity}")
                    i += 1
                # Reset section
                current_section = None
                print(f"--- Exiting section: Extracurricular Activities ---")

            elif current_section == "Volunteer Experience":
                while i < n and lines[i].startswith("- "):
                    volunteer = lines[i][2:].strip()
                    data["volunteer_experience"].append(volunteer)
                    print(f"Appending Volunteer Experience: {volunteer}")
                    i += 1
                # Reset section
                current_section = None
                print(f"--- Exiting section: Volunteer Experience ---")

            elif current_section == "Professional Associations":
                while i < n and lines[i].startswith("- "):
                    association = lines[i][2:].strip()
                    data["professional_associations"].append(association)
                    print(f"Appending Professional Association: {association}")
                    i += 1
                # Reset section
                current_section = None
                print(f"--- Exiting section: Professional Associations ---")

            elif current_section == "References":
                references = []
                while i < n and lines[i] not in sections:
                    if lines[i]:  # Skip empty lines
                        references.append(lines[i])
                        print(f"Appending Reference: {lines[i]}")
                    i += 1
                data["references"] = ' '.join(references)
                print(f"References: {data['references']}")
                # Reset section
                current_section = None
                print(f"--- Exiting section: References ---")

            else:
                print(f"Unknown section: {current_section}. Skipping line {i}.")
                i += 1

        except Exception as e:
            print(f"Error parsing the resume text file at line {i}: {e}")
            i += 1  # Skip to the next line in case of an error

    # Move return statement outside the loop
    print("Finished parsing the resume text file.")
    return data

def set_styles(doc):
    """
    Defines and sets custom styles for the Word document.
    """
    styles = doc.styles

    # Title Style for Name
    if 'NameStyle' not in styles:
        name_style = styles.add_style('NameStyle', WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.name = 'Arial'
        name_style.font.size = Pt(24)
        name_style.font.bold = True
        name_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name_style.paragraph_format.space_after = Pt(12)

    # Heading Style for Sections
    if 'HeadingStyle' not in styles:
        heading_style = styles.add_style('HeadingStyle', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

    # Subheading Style for Work Experience
    if 'SubheadingStyle' not in styles:
        subheading_style = styles.add_style('SubheadingStyle', WD_STYLE_TYPE.PARAGRAPH)
        subheading_style.font.name = 'Arial'
        subheading_style.font.size = Pt(12)
        subheading_style.font.bold = True
        subheading_style.paragraph_format.space_before = Pt(6)
        subheading_style.paragraph_format.space_after = Pt(3)

    # Normal Text Style
    if 'NormalStyle' not in styles:
        normal_style = styles.add_style('NormalStyle', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.line_spacing = Pt(14)
        normal_style.paragraph_format.space_after = Pt(6)

    # Bullet Point Style
    if 'BulletStyle' not in styles:
        bullet_style = styles.add_style('BulletStyle', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = 'Arial'
        bullet_style.font.size = Pt(11)
        bullet_style.paragraph_format.left_indent = Inches(0.25)
        bullet_style.paragraph_format.space_after = Pt(3)

def create_resume_document(data, output_docx):
    """
    Creates a formatted Word document based on the parsed resume data.
    """
    doc = docx.Document()

    # Set default margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Define styles
    set_styles(doc)

    # Add Name
    name_paragraph = doc.add_paragraph(data["name"], style='NameStyle')

    # Add Contact Information
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_paragraph.style = 'NormalStyle'
    contact_text = f"{data['phone']} | {data['email']} | {data['address']}"
    contact_paragraph.add_run(contact_text).italic = True

    # Add spacing after contact info
    doc.add_paragraph()

    # Add Professional Summary
    if data["professional_summary"]:
        doc.add_paragraph("Professional Summary", style='HeadingStyle')
        summary_paragraph = doc.add_paragraph(data["professional_summary"], style='NormalStyle')

    # Add Skills
    if data["skills"]:
        doc.add_paragraph("Skills", style='HeadingStyle')
        for skill in data["skills"]:
            p = doc.add_paragraph(f"• {skill}", style='BulletStyle')

    # Add Certifications and Training
    if data["certifications"]:
        doc.add_paragraph("Certifications and Training", style='HeadingStyle')
        for cert in data["certifications"]:
            p = doc.add_paragraph(f"• {cert}", style='BulletStyle')

    # Add Professional Achievements
    if data["achievements"]:
        doc.add_paragraph("Professional Achievements", style='HeadingStyle')
        for achievement in data["achievements"]:
            p = doc.add_paragraph(f"• {achievement}", style='BulletStyle')

    # Add Work Experience
    if data["work_experience"]:
        doc.add_paragraph("Work Experience", style='HeadingStyle')
        for job in data["work_experience"]:
            # Company Name
            company_paragraph = doc.add_paragraph(job["company"], style='SubheadingStyle')
            # Position
            if job["position"]:
                position_text = job["position"]
                position_paragraph = doc.add_paragraph(position_text, style='NormalStyle')
                position_paragraph.runs[0].bold = True
            # Description
            description_paragraph = doc.add_paragraph(job["description"], style='NormalStyle')

    # Add Education
    if data["education"]:
        doc.add_paragraph("Education", style='HeadingStyle')
        for edu in data["education"]:
            edu_paragraph = doc.add_paragraph(edu, style='NormalStyle')

    # Add Interests
    if data["interests"]:
        doc.add_paragraph("Interests", style='HeadingStyle')
        for interest in data["interests"]:
            p = doc.add_paragraph(f"• {interest}", style='BulletStyle')

    # Add Extracurricular Activities
    if data["extracurricular_activities"]:
        doc.add_paragraph("Extracurricular Activities", style='HeadingStyle')
        for activity in data["extracurricular_activities"]:
            p = doc.add_paragraph(f"• {activity}", style='BulletStyle')

    # Add Volunteer Experience
    if data["volunteer_experience"]:
        doc.add_paragraph("Volunteer Experience", style='HeadingStyle')
        for volunteer in data["volunteer_experience"]:
            p = doc.add_paragraph(f"• {volunteer}", style='BulletStyle')

    # Add Professional Associations
    if data["professional_associations"]:
        doc.add_paragraph("Professional Associations", style='HeadingStyle')
        for association in data["professional_associations"]:
            p = doc.add_paragraph(f"• {association}", style='BulletStyle')

    # Add References
    if data["references"]:
        doc.add_paragraph("References", style='HeadingStyle')
        references_paragraph = doc.add_paragraph(data["references"], style='NormalStyle')

    # Save the document
    doc.save(output_docx)
    print(f"Word document saved as {output_docx}")

def convert_to_pdf(word_file, pdf_file):
    """
    Converts the Word document to PDF using docx2pdf.
    """
    try:
        convert(word_file, pdf_file)
        print(f"PDF document saved as {pdf_file}")
    except Exception as e:
        print(f"Failed to convert to PDF: {e}")

def main():
    # File paths
    text_file = 'tests/generated_tester.txt'          # Input text file
    output_docx = 'resume.docx'       # Output Word document
    output_pdf = 'resume.pdf'         # Output PDF document

    # Parse the text file
    data = parse_resume_text(text_file)

    # Create the Word document
    create_resume_document(data, output_docx)

    # Convert to PDF (optional)
    if os.path.exists(output_docx):
        convert_to_pdf(output_docx, output_pdf)

if __name__ == "__main__":
    main()
